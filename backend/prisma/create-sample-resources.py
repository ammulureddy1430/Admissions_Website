from pathlib import Path
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
PDF_DIR = ROOT / "output" / "pdf"
DOCX_DIR = ROOT / "output" / "docx"
PDF_DIR.mkdir(parents=True, exist_ok=True)
DOCX_DIR.mkdir(parents=True, exist_ok=True)

pdf_path = PDF_DIR / "Ivy-League-SOP-Blueprint-2026.pdf"
styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name="GuideTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=23, leading=28, textColor=colors.HexColor("#071633"), spaceAfter=8))
styles.add(ParagraphStyle(name="GuideSub", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=15, textColor=colors.HexColor("#526474"), spaceAfter=18))
styles.add(ParagraphStyle(name="GuideHead", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=14, leading=18, textColor=colors.HexColor("#008F7D"), spaceBefore=10, spaceAfter=7))
styles.add(ParagraphStyle(name="GuideBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=15, textColor=colors.HexColor("#25364A"), spaceAfter=7))

story = [
    Paragraph("Ivy League SOP Blueprint 2026", styles["GuideTitle"]),
    Paragraph("A practical planning guide for graduate applicants", styles["GuideSub"]),
    Paragraph("The five-part narrative", styles["GuideHead"]),
]
steps = [
    ("1. Opening signal", "Begin with the academic question, experience, or problem that shaped your direction."),
    ("2. Evidence of readiness", "Use two concrete examples showing research, technical, leadership, or analytical preparation."),
    ("3. Program fit", "Connect specific faculty, courses, labs, and communities to your next learning goals."),
    ("4. Future contribution", "Explain the problems you intend to solve and the perspective you will bring."),
    ("5. Closing synthesis", "Return to the opening idea and state a confident, credible next step."),
]
table_data = [[Paragraph("<b>Section</b>", styles["GuideBody"]), Paragraph("<b>Purpose</b>", styles["GuideBody"])]] + [
    [Paragraph(a, styles["GuideBody"]), Paragraph(b, styles["GuideBody"])] for a, b in steps
]
table = Table(table_data, colWidths=[1.55 * inch, 4.65 * inch], repeatRows=1)
table.setStyle(TableStyle([
    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D8EEE8")),
    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#071633")),
    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BCD7CF")),
    ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ("LEFTPADDING", (0, 0), (-1, -1), 8),
    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
    ("TOPPADDING", (0, 0), (-1, -1), 7),
    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
]))
story.extend([
    table,
    Spacer(1, 10),
    Paragraph("Final review checklist", styles["GuideHead"]),
    Paragraph("Use specific nouns and outcomes. Remove generic praise of the university. Verify every faculty and program reference. Keep the statement within the required word limit. Read aloud once for rhythm and clarity.", styles["GuideBody"]),
    Paragraph("<b>Mentor tip:</b> Every paragraph should reveal either evidence, fit, or direction. If it does none of these, revise or remove it.", styles["GuideBody"]),
])
pdf = SimpleDocTemplate(str(pdf_path), pagesize=letter, rightMargin=inch, leftMargin=inch, topMargin=0.8 * inch, bottomMargin=0.8 * inch, title="Ivy League SOP Blueprint 2026")
pdf.build(story)

docx_path = DOCX_DIR / "Professor-LOR-Structure-and-Guidelines.docx"
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.85)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for style_name, size in [("Heading 1", 16), ("Heading 2", 13)]:
    style = doc.styles[style_name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 143, 125)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(4)
run = title.add_run("Professor LOR Structure & Guidelines")
run.bold = True
run.font.name = "Calibri"
run.font.size = Pt(23)
run.font.color.rgb = RGBColor(7, 22, 51)
subtitle = doc.add_paragraph("A concise reference for strong academic recommendations - 2026 admissions cycle")
subtitle.runs[0].font.color.rgb = RGBColor(82, 100, 116)
subtitle.paragraph_format.space_after = Pt(16)

doc.add_heading("Recommended structure", level=1)
sections = [
    ("1. Relationship and context", "State how long and in what capacity you have known the applicant. Name the course, lab, project, or supervision context."),
    ("2. Academic evidence", "Describe one or two specific examples of intellectual ability, research skill, initiative, or disciplined improvement."),
    ("3. Comparative assessment", "Place the student in an honest peer context, such as the top 10% of students taught in the past five years."),
    ("4. Program readiness", "Connect observed strengths to the demands of the proposed degree without making claims outside your direct knowledge."),
    ("5. Clear endorsement", "Close with an unambiguous recommendation and an invitation for follow-up."),
]
for heading, body in sections:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(heading)
    r.bold = True
    r.font.color.rgb = RGBColor(7, 22, 51)
    doc.add_paragraph(body)

doc.add_heading("Quality checklist", level=1)
for item in [
    "Use specific examples rather than adjectives alone.",
    "Keep the letter to one or two pages unless the institution requests otherwise.",
    "Avoid repeating the applicant's statement of purpose or resume.",
    "Use institutional letterhead and include a signature and professional contact details.",
    "Verify program and university names before submission.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Suggested closing", level=1)
closing = doc.add_paragraph()
closing.paragraph_format.left_indent = Inches(0.25)
closing.paragraph_format.right_indent = Inches(0.25)
closing.paragraph_format.space_before = Pt(4)
closing.paragraph_format.space_after = Pt(8)
r = closing.add_run("Based on my direct experience supervising [Applicant Name], I recommend them with confidence for [Program Name]. Their demonstrated strengths in [specific areas] prepare them to contribute meaningfully to your academic community.")
r.italic = True
r.font.color.rgb = RGBColor(37, 54, 74)

doc.save(docx_path)
print(pdf_path)
print(docx_path)
